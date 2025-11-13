import os
import magic  # python-magic for file type detection
import mimetypes
import hashlib
import datetime
from typing import Dict, Any, Optional
import fitz  # PyMuPDF for PDF analysis
from PIL import Image
from ai_analyzer import _call_ollama_understand
import docx  # python-docx for Word documents
import json
import pandas as pd

class FileAnalyzer:
    """Service for analyzing uploaded files and extracting metadata and content."""
    
    def __init__(self):
        self.mime = magic.Magic(mime=True)
    
    def analyze_file(self, file_path: str) -> Dict[str, Any]:
        """Analyze a file and return detailed metadata and content information."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        basic_info = self._get_basic_info(file_path)
        content_info = self._analyze_content(file_path, basic_info["mime_type"])
        
        # Add AI understanding for supported content types
        if content_info.get("extracted_text") or content_info.get("content_summary"):
            try:
                text_to_analyze = content_info.get("extracted_text") or content_info.get("content_summary")
                if text_to_analyze:
                    ai_analysis = _call_ollama_understand(
                        text_to_analyze, 
                        context={
                            "content_type": content_info.get("content_type"),
                            "metadata": content_info.get("metadata", {})
                        }
                    )
                    content_info["ai_understanding"] = ai_analysis
            except Exception as e:
                content_info["ai_understanding"] = {
                    "type": "error",
                    "error": f"AI analysis failed: {str(e)}"
                }
        
        return {
            **basic_info,
            **content_info
        }
    
    def _get_basic_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information like size, type, dates etc."""
        stats = os.stat(file_path)
        file_hash = self._calculate_hash(file_path)
        mime_type = self.mime.from_file(file_path)
        
        return {
            "filename": os.path.basename(file_path),
            "size_bytes": stats.st_size,
            "size_human": self._human_readable_size(stats.st_size),
            "mime_type": mime_type,
            "extension": os.path.splitext(file_path)[1].lower(),
            "created_at": datetime.datetime.fromtimestamp(stats.st_ctime).isoformat(),
            "modified_at": datetime.datetime.fromtimestamp(stats.st_mtime).isoformat(),
            "hash": file_hash,
        }
    
    def _analyze_content(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Analyze file content based on its type."""
        result = {
            "content_type": "unknown",
            "content_summary": None,
            "extracted_text": None,
            "metadata": {}
        }
        
        try:
            # Text files
            if mime_type.startswith("text/"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read(10000)  # Read first 10KB for summary
                    result.update({
                        "content_type": "text",
                        "content_summary": content[:500] + "..." if len(content) > 500 else content,
                        "extracted_text": content,
                        "metadata": {
                            "encoding": "utf-8",
                            "line_count": content.count('\n') + 1
                        }
                    })
            
            # PDF documents
            elif mime_type == "application/pdf":
                doc = fitz.open(file_path)
                text_content = []
                for page in doc:
                    text_content.append(page.get_text())
                full_text = "\n".join(text_content)
                result.update({
                    "content_type": "pdf",
                    "content_summary": full_text[:500] + "..." if len(full_text) > 500 else full_text,
                    "extracted_text": full_text,
                    "metadata": {
                        "page_count": doc.page_count,
                        "pdf_version": doc.pdf_version,
                        "title": doc.metadata.get("title", ""),
                        "author": doc.metadata.get("author", ""),
                        "creation_date": doc.metadata.get("creationDate", "")
                    }
                })
                
            # Images
            elif mime_type.startswith("image/"):
                img = Image.open(file_path)
                result.update({
                    "content_type": "image",
                    "metadata": {
                        "dimensions": f"{img.width}x{img.height}",
                        "format": img.format,
                        "mode": img.mode,
                        "dpi": img.info.get("dpi", None)
                    }
                })
                
            # Word documents
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(file_path)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                result.update({
                    "content_type": "word",
                    "content_summary": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                    "extracted_text": text_content,
                    "metadata": {
                        "paragraph_count": len(doc.paragraphs),
                        "word_count": len(text_content.split())
                    }
                })
                
            # Excel files
            elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                             "application/vnd.ms-excel"]:
                df = pd.read_excel(file_path)
                result.update({
                    "content_type": "spreadsheet",
                    "content_summary": str(df.head()),
                    "metadata": {
                        "row_count": len(df),
                        "column_count": len(df.columns),
                        "columns": list(df.columns)
                    }
                })
                
            # JSON files
            elif mime_type == "application/json" or file_path.endswith('.json'):
                with open(file_path, 'r') as f:
                    data = json.load(f)
                result.update({
                    "content_type": "json",
                    "content_summary": str(data)[:500] + "..." if len(str(data)) > 500 else str(data),
                    "metadata": {
                        "structure": self._analyze_json_structure(data)
                    }
                })
        
        except Exception as e:
            result["error"] = str(e)
            
        return result
    
    def _calculate_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _human_readable_size(self, size_in_bytes: int) -> str:
        """Convert bytes to human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_in_bytes < 1024.0:
                return f"{size_in_bytes:.1f} {unit}"
            size_in_bytes /= 1024.0
        return f"{size_in_bytes:.1f} PB"
    
    def _analyze_json_structure(self, data: Any, max_depth: int = 3) -> Dict[str, Any]:
        """Analyze the structure of a JSON object."""
        if max_depth <= 0:
            return {"type": type(data).__name__}
            
        if isinstance(data, dict):
            return {
                "type": "object",
                "properties": {k: self._analyze_json_structure(v, max_depth - 1) 
                             for k, v in list(data.items())[:10]}  # Limit to first 10 properties
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "sample": [self._analyze_json_structure(x, max_depth - 1) 
                          for x in data[:5]]  # Show first 5 items
            }
        else:
            return {"type": type(data).__name__}